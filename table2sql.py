import docx
import pyodbc
import re


server = '(localdb)\MSSQLLocalDB'
database = 'psdtest3'
conn = pyodbc.connect('DRIVER={SQL Server Native Client 11.0};SERVER='+server+';DATABASE='+database)

doc = docx.Document('tables.docx')
titles = doc.paragraphs
tablesnames = []
for title in titles:
    title = re.findall('[a-zA-Z0-9]+', title.text)
    for i in title:
        tablesnames.append(i)

print(len(tablesnames))


cursor = conn.cursor()
for i in tablesnames:
    cursor.execute("CREATE TABLE {0} (ID int PRIMARY key)".format(i))
    cursor.commit()
cursor.commit()

tables = doc.tables[1:]
print(len(tables))

cursor = conn.cursor()
num = 0
for table in tables:
    for row in table.rows[1:]:
        t1 = row.cells[0].text
        t2 = row.cells[1].text
        if t1 == "ID":
            current_table = tablesnames[num]
            print("----------" + current_table + "----------")
            num += 1
        else:
            print(t1 + "\t" + t2)
            cursor.execute("ALTER table {0} add {1} {2}".format(current_table, t1, t2))
            cursor.commit()
            # print("ALTER table {0} add {1} {2}".format(current_table, t1, t2))

# cursor.commit()
print("done")
print(num)



